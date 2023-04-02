from ..document import Document
from .block import *
from docx import Document as docxDocument

from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table, _Row
from docx.text.paragraph import Paragraph


class WordDocument(Document):
    def __init__(self, path):
        super().__init__(path)

    def _parse_fields(self):
        def get_table_type(docx_table):
            if len(docx_table.columns) == 2:
                return "simple_form"
            else:
                raise NotImplementedError("Not ready to work with different tables!")
                return "smth other"

        def iter_block_items(parent):
            """
            Generate a reference to each paragraph and table child within *parent*,
            in document order. Each returned value is an instance of either Table or
            Paragraph. *parent* would most commonly be a reference to a main
            Document object, but also works for a _Cell object, which itself can
            contain paragraphs and tables.
            """
            if isinstance(parent, _Document):
                parent_elm = parent.element.body
            elif isinstance(parent, _Cell):
                parent_elm = parent._tc
            elif isinstance(parent, _Row):
                parent_elm = parent._tr
            else:
                raise ValueError("something's not right")
            for child in parent_elm.iterchildren():
                if isinstance(child, CT_P):
                    yield Paragraph(child, parent)
                elif isinstance(child, CT_Tbl):
                    yield Table(child, parent)

        fields = {}
        n_paragraphs = 0
        n_tables = 0

        for block in iter_block_items(self._f_instance):
            if isinstance(block, Paragraph):
                # TODO: ignore empty paragraphs?
                fields[f"paragraph_{n_paragraphs}"] = WordDocumentTextBlock(block, f"paragraph_{n_paragraphs}")
                n_paragraphs += 1
            elif isinstance(block, Table):
                cells = {}
                if get_table_type(block) == "simple_form":
                    for row in block.rows:
                        field_name = row.cells[0].text
                        cells[field_name] = WordDocumentTableField(row.cells[1],
                                                                   field_name)
                    fields[f"table_{n_tables}"] = WordDocumentTable(block,
                                                                    f"table_{n_tables}",
                                                                    cells)
                else:
                    raise NotImplementedError("Other tables are not supported yet!")

                n_tables += 1

        return fields

    def insert_row(self, table_name, field_name):
        if table_name not in self._fields:
            raise Exception(f"Table {table_name} doesn't exist!")

        table_num = table_name.removeprefix("table_")
        f_table: Table = self._f_instance.tables[int(table_num)]

        new_row = f_table.add_row()
        new_row.cells[0].text = field_name

        table_field = WordDocumentTableField(new_row.cells[1], field_name)
        table: WordDocumentTable = self._fields[table_name]
        table.add_field(table_field, field_name)

    def remove_row(self, table_name, field_name):
        if table_name not in self._fields:
            raise Exception(f"Table \"{table_name}\" doesn't exist!")

        table_num = table_name.removeprefix("table_")

        f_table = self._f_instance.tables[int(table_num)]

        row = next(x for x in f_table.rows if x.cells[0].text == field_name)

        tbl = f_table._tbl
        tr = row._tr
        tbl.remove(tr)

        table: WordDocumentTable = self._fields[table_name]
        table.remove_field(field_name)

    def open(self):
        return docxDocument(self._path)

    def save(self, path):
        self._f_instance.save(path)
